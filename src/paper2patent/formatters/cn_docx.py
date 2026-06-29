"""CN patent .docx writer — assembles PatentIR into a formatted .docx file."""

from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT

from paper2patent.ir import PatentIR, StyleProfile
from paper2patent.formatters.styles import apply_styles_to_docx


class CNDocxWriter:
    """Write a PatentIR to a CN-formatted .docx file."""

    def __init__(self, output_dir: str, style_profile: StyleProfile):
        self.output_dir = Path(output_dir)
        self.style = style_profile

    def write(self, patent: PatentIR) -> Path:
        """Write the full CN patent .docx and return the output path."""
        doc = Document()

        # Apply style profile
        self._setup_page(doc)

        # ==================================================================
        # Page 1: 说明书摘要 + 摘要附图
        # ==================================================================
        self._write_abstract(doc, patent)

        # ==================================================================
        # Page 2: 权利要求书
        # ==================================================================
        doc.add_page_break()
        self._write_claims(doc, patent)

        # ==================================================================
        # Page 3+: 说明书 (5 mandatory sections)
        # ==================================================================
        doc.add_page_break()
        self._write_title(doc, patent)
        self._write_sections(doc, patent)

        # ==================================================================
        # Drawings section (说明书附图)
        # ==================================================================
        self._write_drawings(doc, patent)

        # Save
        safe_title = self._sanitize_filename(patent.title or "专利")
        output_path = self.output_dir / f"{safe_title}.docx"
        doc.save(str(output_path))

        return output_path

    # ------------------------------------------------------------------
    # Page setup
    # ------------------------------------------------------------------

    def _setup_page(self, doc: Document) -> None:
        """Apply CN patent page settings."""
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(self.style.margin_top_mm / 10)
            section.bottom_margin = Cm(self.style.margin_bottom_mm / 10)
            section.left_margin = Cm(self.style.margin_left_mm / 10)
            section.right_margin = Cm(self.style.margin_right_mm / 10)

        # Normal style
        style = doc.styles["Normal"]
        style.font.size = Pt(self.style.body_font_size_pt)
        style.font.name = self.style.body_font_family
        style.paragraph_format.line_spacing = self.style.line_spacing

    # ------------------------------------------------------------------
    # Abstract
    # ------------------------------------------------------------------

    def _write_abstract(self, doc: Document, patent: PatentIR) -> None:
        """Write 说明书摘要."""
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("说明书摘要")
        title_run.bold = True
        title_run.font.size = Pt(self.style.heading_font_size_pt)
        title_run.font.name = self.style.heading_font_family

        doc.add_paragraph("")  # blank line

        # Abstract content
        abstract_text = patent.abstract or "(摘要内容)"
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(
            self.style.body_font_size_pt * 0.035 * self.style.first_line_indent_chars
        )
        run = para.add_run(abstract_text)
        run.font.name = self.style.abstract_font_family
        run.font.size = Pt(self.style.abstract_font_size_pt)

        # Abstract figure (if available)
        if patent.abstract_figure_path and os.path.exists(patent.abstract_figure_path):
            doc.add_paragraph("")
            fig_para = doc.add_paragraph()
            fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                fig_para.add_run().add_picture(
                    patent.abstract_figure_path,
                    width=Inches(2.5),
                )
            except Exception:
                fig_para.add_run(f"[附图: {patent.abstract_figure_path}]")

    # ------------------------------------------------------------------
    # Claims
    # ------------------------------------------------------------------

    def _write_claims(self, doc: Document, patent: PatentIR) -> None:
        """Write 权利要求书."""
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("权利要求书")
        title_run.bold = True
        title_run.font.size = Pt(self.style.heading_font_size_pt)
        title_run.font.name = self.style.heading_font_family

        doc.add_paragraph("")

        for claim in patent.claims:
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(
                self.style.body_font_size_pt * 0.035 * self.style.first_line_indent_chars
            )

            # Number
            num_run = para.add_run(f"{claim.number}. ")
            num_run.font.name = self.style.claims_font_family
            num_run.font.size = Pt(self.style.claims_font_size_pt)

            # Claim text
            text_run = para.add_run(claim.text)
            text_run.font.name = self.style.claims_font_family
            text_run.font.size = Pt(self.style.claims_font_size_pt)

            doc.add_paragraph("")  # spacing between claims

    # ------------------------------------------------------------------
    # Description — title
    # ------------------------------------------------------------------

    def _write_title(self, doc: Document, patent: PatentIR) -> None:
        """Write invention title, centered at top of 说明书."""
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(patent.title or "(发明名称)")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = self.style.heading_font_family

        doc.add_paragraph("")  # blank line after title

    # ------------------------------------------------------------------
    # Description — sections
    # ------------------------------------------------------------------

    def _write_sections(self, doc: Document, patent: PatentIR) -> None:
        """Write all 【】sections."""
        for section in patent.sections:
            self._write_section_heading(doc, section.heading)
            self._write_section_content(doc, section.content)

    def _write_section_heading(self, doc: Document, heading: str) -> None:
        """Write a 【section heading】 in bold."""
        para = doc.add_paragraph()
        run = para.add_run(heading)
        run.bold = True if self.style.heading_bold else False
        run.font.size = Pt(self.style.heading_font_size_pt)
        run.font.name = self.style.heading_font_family
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

    def _write_section_content(self, doc: Document, content: str) -> None:
        """Write section body content with basic formatting."""
        paragraphs = content.split("\n")

        for text in paragraphs:
            text = text.strip()
            if not text:
                continue

            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(
                self.style.body_font_size_pt * 0.035 * self.style.first_line_indent_chars
            )

            # Handle basic markdown in text
            self._add_formatted_run(para, text)

    def _add_formatted_run(self, para, text: str) -> None:
        """Add a paragraph run with basic bold/italic handling."""
        # Bold markers: **text**
        parts = re.split(r'(\*\*(.+?)\*\*)', text)
        i = 0
        while i < len(parts):
            part = parts[i]
            if i + 2 < len(parts) and part.startswith("**"):
                # Bold
                run = para.add_run(parts[i + 1])
                run.bold = True
                run.font.name = self.style.body_font_family
                run.font.size = Pt(self.style.body_font_size_pt)
                i += 3
            else:
                run = para.add_run(part)
                run.font.name = self.style.body_font_family
                run.font.size = Pt(self.style.body_font_size_pt)
                i += 1

    # ------------------------------------------------------------------
    # Drawings
    # ------------------------------------------------------------------

    def _write_drawings(self, doc: Document, patent: PatentIR) -> None:
        """Write 说明书附图 — one figure per page."""
        all_figures = patent.figures + patent.diagrams

        if not all_figures:
            return

        for fig_info in all_figures:
            doc.add_page_break()

            fig_path = fig_info.get("path", "")
            caption = fig_info.get("caption", "")

            # Caption above or below figure
            if caption:
                cap_para = doc.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap_para.add_run(caption)
                cap_run.font.name = self.style.body_font_family
                cap_run.font.size = Pt(10)

            # Insert figure
            if fig_path and os.path.exists(fig_path):
                fig_para = doc.add_paragraph()
                fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    fig_para.add_run().add_picture(
                        fig_path,
                        width=Inches(5.5),
                    )
                except Exception:
                    fig_para.add_run(f"[图片: {os.path.basename(fig_path)}]")
            else:
                doc.add_paragraph(f"[{caption}]")

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _sanitize_filename(self, name: str) -> str:
        """Remove illegal filename characters."""
        illegal = r'[<>:"/\\|?*]'
        safe = re.sub(illegal, "_", name)
        return safe[:80]  # limit length

    # ------------------------------------------------------------------
    # Format Preview Document
    # ------------------------------------------------------------------

    def write_format_preview(self) -> Path:
        """Generate a standalone .docx showing current patent format settings.

        This document serves as both a format specification and visual preview,
        so the user can verify fonts, margins, spacing before full conversion.
        """
        doc = Document()
        self._setup_page(doc)

        # ================================================================
        # Cover / Title
        # ================================================================
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("专利格式预览 — Patent Format Preview")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.name = self.style.heading_font_family

        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle_para.add_run(
            f"专利类型: {self.style.name}  |  "
            f"样式来源: {self.style.source}  |  "
            f"纸张: {self.style.page_size}"
        )
        sub_run.font.size = Pt(10)
        sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")

        # ================================================================
        # Section 1: 页面设置 (Page Settings)
        # ================================================================
        self._preview_section_heading(doc, "一、页面设置 (Page Setup)")

        settings_table = doc.add_table(rows=00, cols=2)
        settings_table.style = "Light Grid Accent 1"
        self._add_table_row(settings_table, "纸张大小", self.style.page_size)
        self._add_table_row(settings_table, "上边距", f"{self.style.margin_top_mm} mm")
        self._add_table_row(settings_table, "下边距", f"{self.style.margin_bottom_mm} mm")
        self._add_table_row(settings_table, "左边距", f"{self.style.margin_left_mm} mm")
        self._add_table_row(settings_table, "右边距", f"{self.style.margin_right_mm} mm")

        doc.add_paragraph("")

        # ================================================================
        # Section 2: 字体与排版 (Font & Typography)
        # ================================================================
        self._preview_section_heading(doc, "二、字体与排版 (Font & Typography)")

        font_table = doc.add_table(rows=00, cols=3)
        font_table.style = "Light Grid Accent 1"
        self._add_table_row(font_table, "元素", "字体", "字号")
        self._add_table_row(font_table, "正文", self.style.body_font_family, f"{self.style.body_font_size_pt} pt")
        self._add_table_row(font_table, "章节标题", self.style.heading_font_family, f"{self.style.heading_font_size_pt} pt")
        self._add_table_row(font_table, "权利要求", self.style.claims_font_family, f"{self.style.claims_font_size_pt} pt")
        self._add_table_row(font_table, "摘要", self.style.abstract_font_family, f"{self.style.abstract_font_size_pt} pt")

        doc.add_paragraph("")

        # ================================================================
        # Section 3: 段落格式 (Paragraph Format)
        # ================================================================
        self._preview_section_heading(doc, "三、段落格式 (Paragraph Format)")

        para_table = doc.add_table(rows=00, cols=2)
        para_table.style = "Light Grid Accent 1"
        self._add_table_row(para_table, "行距", f"{self.style.line_spacing} 倍")
        self._add_table_row(para_table, "首行缩进", f"{self.style.first_line_indent_chars} 字符")
        self._add_table_row(para_table, "章节标题括号", self.style.section_title_brackets)

        doc.add_paragraph("")

        # ================================================================
        # Section 4: 视觉样例 (Visual Samples)
        # ================================================================
        self._preview_section_heading(doc, "四、视觉样例 (Visual Samples)")

        # --- Sample: Invention Title ---
        self._preview_section_heading(doc, "▎发明名称 (Title)")
        sample_title = doc.add_paragraph()
        sample_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = sample_title.add_run("一种基于深度学习的图像处理方法及装置")
        tr.bold = True
        tr.font.size = Pt(16)
        tr.font.name = self.style.heading_font_family

        _ = doc.add_paragraph("")  # blank line after title per CNIPA

        # --- Sample: Section Heading ---
        self._preview_section_heading(doc, "▎章节标题 (Section Heading)")
        sample_heading = doc.add_paragraph()
        hr = sample_heading.add_run("【技术领域】")
        hr.bold = True
        hr.font.size = Pt(self.style.heading_font_size_pt)
        hr.font.name = self.style.heading_font_family

        # --- Sample: Body Text ---
        self._preview_section_heading(doc, "▎正文段落 (Body Paragraph)")
        body_para = doc.add_paragraph()
        body_para.paragraph_format.first_line_indent = Cm(
            self.style.body_font_size_pt * 0.035 * self.style.first_line_indent_chars
        )
        body_text = (
            "本发明涉及人工智能技术领域，具体涉及一种图像处理方法及装置。"
            "随着深度学习技术的快速发展，卷积神经网络在图像识别、目标检测等任务中取得了显著成果。"
            "然而，现有技术中的深度神经网络模型通常需要大量的计算资源和存储空间，"
            "难以部署在资源受限的边缘设备上。"
        )
        br = body_para.add_run(body_text)
        br.font.name = self.style.body_font_family
        br.font.size = Pt(self.style.body_font_size_pt)

        # Second paragraph to show spacing
        body_para2 = doc.add_paragraph()
        body_para2.paragraph_format.first_line_indent = Cm(
            self.style.body_font_size_pt * 0.035 * self.style.first_line_indent_chars
        )
        body_text2 = (
            "为解决上述技术问题，本发明提供一种模型压缩方法，通过知识蒸馏和剪枝技术的结合，"
            "在保持模型精度的同时显著减小模型体积。以下结合附图和实施例对本发明进行详细说明。"
        )
        br2 = body_para2.add_run(body_text2)
        br2.font.name = self.style.body_font_family
        br2.font.size = Pt(self.style.body_font_size_pt)

        # --- Sample: Claims ---
        self._preview_section_heading(doc, "▎权利要求 (Claims)")

        for claim_num, claim_text in [
            (1, "一种图像处理方法，其特征在于，包括：\n步骤S110：获取输入图像；\n步骤S120：对输入图像进行特征提取，得到特征图；\n步骤S130：对特征图进行分类，输出分类结果。"),
            (2, "根据权利要求1所述的方法，其特征在于，所述特征提取使用卷积神经网络。"),
            (3, "根据权利要求1所述的方法，其特征在于，所述分类使用Softmax分类器。"),
        ]:
            cp = doc.add_paragraph()
            cp.paragraph_format.first_line_indent = Cm(
                self.style.body_font_size_pt * 0.035 * self.style.first_line_indent_chars
            )
            cn = cp.add_run(f"{claim_num}. ")
            cn.font.name = self.style.claims_font_family
            cn.font.size = Pt(self.style.claims_font_size_pt)
            ct = cp.add_run(claim_text)
            ct.font.name = self.style.claims_font_family
            ct.font.size = Pt(self.style.claims_font_size_pt)

        # --- Sample: Abstract ---
        self._preview_section_heading(doc, "▎摘要 (Abstract)")

        abs_para = doc.add_paragraph()
        abs_para.paragraph_format.first_line_indent = Cm(
            self.style.body_font_size_pt * 0.035 * self.style.first_line_indent_chars
        )
        abs_text = (
            "本发明提供一种图像处理方法及装置。该方法包括：获取输入图像；"
            "对输入图像进行特征提取；对提取的特征进行分类，输出分类结果。"
            "本发明通过轻量级网络结构设计，在保持高精度的同时降低了计算复杂度，"
            "适用于边缘设备部署。"
        )
        ar = abs_para.add_run(abs_text)
        ar.font.name = self.style.abstract_font_family
        ar.font.size = Pt(self.style.abstract_font_size_pt)

        # --- Sample: Figure Caption ---
        self._preview_section_heading(doc, "▎附图说明 (Figure Caption)")
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run("图1 是本发明实施例提供的图像处理方法的流程示意图")
        cap_run.font.name = self.style.body_font_family
        cap_run.font.size = Pt(10)

        # ================================================================
        # Section 5: Notes
        # ================================================================
        doc.add_page_break()
        self._preview_section_heading(doc, "五、说明 (Notes)")

        notes = [
            "• 以上为当前样式配置下的专利格式预览。",
            "• 如需修改格式，可通过以下方式：",
            "   1. 使用 -r 参数指定参考专利 .docx 文件以提取其样式",
            "   2. 编辑 src/paper2patent/formatters/default_cn_styles.json 修改内置默认值",
            "   3. 在转换完成后手动调整输出 .docx",
            "• 实际转换时，页面设置（边距、纸张）将应用于文档所有页面。",
            "• 章节标题使用【】括号格式是 CNIPA 标准要求。",
            "• 附图中的引用号需在说明书中全部提及，系统会在 Step 9 自动进行交叉引用检查。",
        ]
        for note in notes:
            np = doc.add_paragraph()
            if note.startswith("   "):
                nr = np.add_run(note)
                nr.font.size = Pt(9)
            else:
                nr = np.add_run(note)
                nr.font.size = Pt(10)
            nr.font.name = self.style.body_font_family

        # Save
        output_path = self.output_dir / "专利格式预览.docx"
        doc.save(str(output_path))
        return output_path

    # ------------------------------------------------------------------
    # Format preview helpers
    # ------------------------------------------------------------------

    def _preview_section_heading(self, doc: Document, heading: str) -> None:
        """Write a format preview section heading."""
        para = doc.add_paragraph()
        run = para.add_run(heading)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = self.style.heading_font_family
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(4)

    def _add_table_row(self, table, *cells: str) -> None:
        """Add a row to a format preview table."""
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            row.cells[i].text = cell_text
            for para in row.cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = self.style.body_font_family
