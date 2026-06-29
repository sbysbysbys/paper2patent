"""Extract style information from a reference .docx patent."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, Cm

from paper2patent.ir import StyleProfile


class ReferenceReader:
    """Read a reference .docx patent and extract its formatting styles."""

    def extract(self, docx_path: str) -> StyleProfile:
        """Extract a StyleProfile from an existing .docx patent file."""
        doc = Document(docx_path)

        profile = StyleProfile(
            name="reference",
            source=f"reference:{docx_path}",
        )

        # --- Page margins ---
        if doc.sections:
            section = doc.sections[0]
            profile.margin_top_mm = self._cm_to_mm(section.top_margin)
            profile.margin_bottom_mm = self._cm_to_mm(section.bottom_margin)
            profile.margin_left_mm = self._cm_to_mm(section.left_margin)
            profile.margin_right_mm = self._cm_to_mm(section.right_margin)

            # Page size
            if section.page_width and section.page_height:
                w_mm = section.page_width / 360000 * 25.4  # EMU -> mm
                h_mm = section.page_height / 360000 * 25.4
                if abs(w_mm - 210) < 5 and abs(h_mm - 297) < 5:
                    profile.page_size = "A4"
                elif abs(w_mm - 215.9) < 5:
                    profile.page_size = "Letter"

        # --- Body font (from Normal style) ---
        normal = doc.styles["Normal"]
        profile.body_font_family = normal.font.name or profile.body_font_family
        if normal.font.size:
            profile.body_font_size_pt = normal.font.size.pt
        if normal.font.color and normal.font.color.rgb:
            profile.body_font_color = str(normal.font.color.rgb)

        # Line spacing
        if normal.paragraph_format.line_spacing:
            profile.line_spacing = float(normal.paragraph_format.line_spacing)

        # --- Heading font (from Heading 1 style) ---
        try:
            h1 = doc.styles["Heading 1"]
            profile.heading_font_family = h1.font.name or profile.heading_font_family
            if h1.font.size:
                profile.heading_font_size_pt = h1.font.size.pt
            profile.heading_bold = h1.font.bold if h1.font.bold is not None else True
        except KeyError:
            pass  # Use defaults

        # --- Detect Chinese font usage ---
        # Scan first few paragraphs to check for Chinese fonts
        cn_fonts_found = set()
        for para in doc.paragraphs[:20]:
            for run in para.runs[:5]:
                fn = run.font.name
                if fn and any(cn in fn for cn in ("宋", "黑", "仿宋", "楷", "隶")):
                    cn_fonts_found.add(fn)

        if cn_fonts_found:
            # If Chinese fonts detected, use them
            sorted_fonts = sorted(cn_fonts_found)
            if any("宋" in f for f in sorted_fonts):
                profile.body_font_family = next(f for f in sorted_fonts if "宋" in f)

        # --- Detect section title pattern ---
        for para in doc.paragraphs[:30]:
            text = para.text.strip()
            if text.startswith("【") and "】" in text:
                profile.section_title_brackets = "【】"
                break
            elif text.startswith("[") and "]" in text[:20]:
                profile.section_title_brackets = "[]"
                break

        return profile

    @staticmethod
    def _cm_to_mm(value) -> float:
        """Convert docx length units to mm."""
        if value is None:
            return 0.0
        try:
            return round(value / 360000 * 25.4, 1)  # EMU -> mm
        except Exception:
            return float(value) if value else 0.0
